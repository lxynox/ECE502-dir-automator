from docx import Document
from docx.shared import Inches
import sys
import os

document = Document()
# Title section
document.add_heading('Assignment #', 0)
# Paragraphs & Font styles
p = document.add_paragraph('Below is the ')
p.add_run('Solution').bold = True
p.add_run(' to the ')
p.add_run('problem.').italic = True
# Heading
document.add_heading('Solution', level=1)
document.add_paragraph('insert your picture below', style='IntenseQuote')

# document.add_paragraph(
#     'first item in unordered list', style='ListBullet'
# )
# document.add_paragraph(
#     'first item in ordered list', style='ListNumber'
# )

# Picture
# TODO: remove the following test scripts
dirname, filename = os.path.split(os.path.abspath(__file__))
print 
print "current working directory, " + os.getcwd()
print "running from", dirname
print "file is", filename
print "command line arguments: " + str(sys.argv) 

# support picture format both png & jpeg 
document.add_picture(sys.argv[1]+"/images/"+ sys.argv[2], width=Inches(5))
#document.add_picture(sys.argv[1]+"/images/"+ sys.argv[2] + #".jpeg", width=Inches(5))

# table = document.add_table(rows=1, cols=3)
# hdr_cells = table.rows[0].cells
# hdr_cells[0].text = 'Qty'
# hdr_cells[1].text = 'Id'
# hdr_cells[2].text = 'Desc'
"""
for item in recordset:
    row_cells = table.add_row().cells
    row_cells[0].text = str(item.qty)
    row_cells[1].text = str(item.id)
    row_cells[2].text = item.desc
"""
document.add_page_break()
# Name and Save it
document.save("Problem" + sys.argv[2][:sys.argv[2].index(".")] + ".docx")
